# -*- coding: utf-8 -*-
"""
Created on Fri Mar 27 22:27:24 2020

@author: Chenglong
"""
"""
环境要求：python 3.7, python-docx 模块
环境配置方法：https://www.python.org/downloads/
https://phoenixnap.com/kb/how-to-install-python-3-windows
https://blog.csdn.net/qq_38628350/article/details/79155491
python-docx 官方文档：https://python-docx.readthedocs.io/en/latest/
fileGadget主要实现了从指定docx文件读取指定段落并拷贝到指定文件功能
"""
"""
版本： Beta 1.0.0
日期： Fri Mar 27 22:27:24 2020
版本更改描述：完成了基本需求- 1.从目标文件读取内容；2.获取搜索关键词；3.使用获取的关键词到源文件中查找内容；4.使用关键词确定位置，插入内容；
命名规范：下划线命名法

版本： Beta 2.0.0
日期： Sun Apr 5 17:24:00 2020
版本更改描述：修复了表格与文本对象无法复制到新文档的问题，目前问题是找不到项目基本情况表对应位置，剩余解决是删除索引段落；
命名规范：下划线命名法

版本： Beta 2.1.0
日期： Sun Apr 5 17:43:00 2020
版本更改描述：修复了表格与文本对象无法复制到新文档的问题，删除了不需要的索引段落，目前问题是找不到项目基本情况表对应位置；
命名规范：下划线命名法

版本： Beta 2.2.0
日期： Sun Apr 5 17:52:00 2020
版本更改描述：修复了表格与文本对象无法复制到新文档的问题，删除了不需要的索引段落，移除了不必要的括号，目前问题是找不到项目基本情况表对应位置，并且拷贝内容有多余的标题，缺少原始的序号；
命名规范：下划线命名法

"""

"""
 导包--
 docx 包结构 https://pydoc.net/python-docx/0.8.6/
"""

from docx import Document
from win32com import client as wc
import os
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.section import CT_PageMar
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from copy import deepcopy
import re

#获取当前路径
current_dir = os.path.dirname(os.path.abspath(__file__))
#将 doc 文件转化为 docx 文件，路径必须为全路径！
word = wc.Dispatch("Word.Application")
org_target = 'XXX股份有限公司IPO项目风险评级报告——【】级 1.doc'
org_source = '1、IPO项目立项申请报告.doc'
org_target_full = os.path.join(current_dir, org_target)
org_source_full = os.path.join(current_dir, org_source)

doc_target = word.Documents.Open(org_target_full)
doc_target.SaveAs(org_target_full+'x', 16)
doc_target.Close()

doc_source = word.Documents.Open(org_source_full)
doc_source.SaveAs(org_source_full+'x', 16)
doc_source.Close()

word.Quit()


#读取目标文件
target_filename = org_target_full+'x'
source_filename = org_source_full+'x'

target_document = Document(target_filename)
para_count = 0
query_list = []
pos_list = []
content_list = []
# 返回表格中的目标单元格内容


def table_reader(Tableg, query_array):
    result_list = []
    row_cells, column_cells = [], []
    index = []
    width, length = len(Tableg.columns), len(Tableg.rows)
    k = 0
    for row in Tableg.rows:
        for cell in row.cells:
            if cell not in row_cells:
                index.append([k // width, k % width])
                row_cells.append(cell)
            k += 1
    k = 0
    for column in Tableg.columns:
        for cell in column.cells:
            if cell not in column_cells:
                column_cells.append(cell)
            elif [k % length, k // length] in index:
                index.remove([k % length, k // length])
            k += 1
    # index即为找到的单元格索引
    #     print(index)
    for i in index:
        if query_array in Tableg.rows[i[0]].cells[i[1]].text:
            result_list.append(Tableg.rows[i[0]].cells[i[1] + 1].text)
    return result_list


# 读取目标文件
for para in target_document.paragraphs:
    if '【立项申请报告' in para.text:
        pos_start = para.text.find('告')
        pos_end = para.text.find('】')
        # print(para.text[pos_start + 1:pos_end])
        query_list.append(para.text[pos_start + 1:pos_end])
        pos_list.append(para_count)
        # print(pos_list)
        # p = para._element
        # p.getparent().remove(p)
        # p._p = p._element = None
    para_count += 1

# 打开源文件
source_document = Document(source_filename)
# print('\n')
sets = source_document.element.body.xpath('w:p | w:tbl')
# print(query_list)
for pointer in range(len(query_list)):
    # print('this query:', query_list[pointer].split('、')[1])
    for i in range(len(sets)):
        if isinstance(sets[i], CT_P):
            para = Paragraph(sets[i], source_document)
            if query_list[pointer].split('、')[1] in para.text:
                # print('Found:')
                # 获取接下来迭代的内容
                temp_list = []
                for j in range(i+1, len(sets)):
                    # 遇到标题则break
                    if isinstance(sets[j], CT_P):
                        paragraph_temp = Paragraph(sets[j], source_document)
                        paragraph_temp.text = paragraph_temp.text.replace('【', '').replace('】', '')
                        # print(paragraph_temp.text)
                        # print('------------------->')
                        temp_list.append(paragraph_temp)
                        # if re.findall('\([一|二|三|四|五|六|七|八|九|十]*\)', paragraph_temp.text) or re.findall('[一|二|三|四|五|六|七|八|九|十]*、', paragraph_temp.text):
                        #     break
                        # if re.findall('[一|二|三|四|五|六|七|八|九|十]*、', paragraph_temp.text):
                        if re.findall("\（[一|二|三|四|五|六|七|八|九|十]\）", paragraph_temp.text) or re.findall('[一|二|三|四|五|六|七|八|九|十]\、', paragraph_temp.text):
                            del (temp_list[-1])
                            break
                    if isinstance(sets[j], CT_Tbl):
                        table_temp = Table(sets[j], source_document)
                        temp_list.append(table_temp)
                        # del (temp_list[-1])
                # print(len(temp_list), 'elements')
                # del(temp_list[-1])
                content_list.append(temp_list)
                break
            # if '项目基本情况表' in query_list[pointer]:
            # if query_list[pointer].split('、')[1] not in para.text:
            #     temp_list = []
            #     content_list.append(temp_list)
            #     break
            if '项目基本情况表' in query_list[pointer] and '项目基本情况表' in para.text:
                # print(query_list[pointer].split('、')[1])
                # print('Found tb:')
                temp_list = []
                abc_list = []
                for j in range(i+1, len(sets)):
                    if isinstance(sets[j], CT_P):
                        paragraph_temp = Paragraph(sets[j], source_document)
                        # temp_list.append(paragraph_temp)
                        # print(paragraph_temp.text)
                        # print('------------------->')
                        if re.findall('\([一|二|三|四|五|六|七|八|九|十]\)', paragraph_temp.text) or re.findall(
                                '[一|二|三|四|五|六|七|八|九|十]\、', paragraph_temp.text):
                            break
                    if isinstance(sets[j], CT_Tbl):
                        table_temp = Table(sets[j], source_document)
                        # temp_list.append(table_temp)
                        # temp_list.append(bcd_list)
                        # content_list.append(temp_list)
                # print(len(temp_list), 'tables')
                # print(temp_list)
                        Search = ['所属行业', '主营业务', '其他中介机构情况', '主要财务数据（万元）']
                        for t in range(len(Search)):
                            if Search[t] in query_list[pointer]:
                                # print(table_reader(table_temp, Search[t]))
                                abc_list.append(table_reader(table_temp, Search[t]))
                                content_list.append(abc_list)
                break

for it in range(len(query_list)):
    # print(query_list[it])
    for parag in target_document.paragraphs:
        if query_list[it] in parag.text and len(content_list[it]) > 0:
            content_list[it].reverse()
            # del(content_list[it][0])
            # print(content_list[it])
            for obj in content_list[it]:
                if isinstance(obj, Table):
                    # tableq = target_document.add_table(len(obj.rows),len(obj.columns))
                    tbl = obj.tbl
                    new_tbl = deepcopy(tbl)
                    parag._p.addnext(new_tbl)
                if isinstance(obj, Paragraph):
                    # paragraphq = target_document.add_paragraph(obj.text)
                    parat = obj.p
                    new_para = deepcopy(parat)
                    parag._p.addnext(new_para)
                if isinstance(obj, list):
                    # print(obj[0])
                    # print(parag.text)
                    paragtext = deepcopy(obj[0])
                    parag.text = paragtext
            target_document.save('abc.docx')
            break

# 移除原索引
for para in target_document.paragraphs:
    if '【立项申请报告' in para.text:
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None

target_document.save('abc.docx')

